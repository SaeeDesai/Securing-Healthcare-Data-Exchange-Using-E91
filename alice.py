from io import BytesIO
import sys
import tkinter as tk
from functools import partial
import hashlib
from tkinter import filedialog
from cryptography.hazmat.primitives import padding
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.backends import default_backend
import socket
import os
import time
import hmac
from datetime import datetime
import PyPDF2
from docx import Document

blocked_ips = []

# Define allowed IP addresses for nurses and doctors
allowed_nurse_ips = ['127.0.0.3']  # Example nurse IP addresses
allowed_doctor_ips = ['127.0.0.1','127.0.0.2']  # Example doctor IP addresses

# Read AES key from the file
with open("final_key_alice.txt", "rb") as key_file:
    aes_key = key_file.read(16)

def delete_key_file():
    try:
        os.remove("final_key_alice.txt")
        print("AES key file deleted.")
    except FileNotFoundError:
        pass  # Ignore if the file doesn't exist

def generate_mac_id_from_key(message, key):
    mac_id = hmac.new(key, message.encode(), hashlib.sha256).hexdigest()
    return mac_id

# Function to read the content of a Word file
def read_word_file(file_path):
    doc = Document(file_path)
    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return content

# Function to read the content of a PDF file
def read_pdf_file(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfFileReader(file)
        content = ""
        for page in range(reader.numPages):
            content += reader.getPage(page).extractText()
    return content

# Function to write content to a Word file
def write_word_file(content, file_path):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)

# Function to write content to a PDF file
def write_pdf_file(content, file_path):
    with open(file_path, 'wb') as file:
        writer = PyPDF2.PdfFileWriter()
        writer.addPage(PyPDF2.PdfFileReader(BytesIO(content)).getPage(0))
        writer.write(file)

def decrypt_message(encrypted_message, key):
    start_time = time.time()
    iv = encrypted_message[:16]  # Extract IV
    cipher = Cipher(algorithms.AES(key), modes.CBC(iv))
    decryptor = cipher.decryptor()
    unpadder = padding.PKCS7(algorithms.AES.block_size).unpadder()

    decrypted_data = decryptor.update(encrypted_message[16:]) + decryptor.finalize()
    unpadded_data = unpadder.update(decrypted_data) + unpadder.finalize()
    decryption_time = time.time() - start_time
    print(f"Decryption time for message: {decryption_time:.2f} s")

    return unpadded_data.decode('utf-8', errors='replace').strip()


def decrypt_file(encrypted_content, key, file_path):
    start_time = time.time()
    try:
        iv = encrypted_content[:16]  # Extract IV

        cipher = Cipher(algorithms.AES(key), modes.CBC(iv), backend=default_backend())
        decryptor = cipher.decryptor()
        unpadder = padding.PKCS7(algorithms.AES.block_size).unpadder()
        decrypted_data = decryptor.update(encrypted_content[16:]) + decryptor.finalize()
        unpadded_data = unpadder.update(decrypted_data) + unpadder.finalize()
        with open(file_path[:-10], 'wb') as decrypted_file:  # Remove '.encrypted' from file name
            decrypted_file.write(unpadded_data)
        print("File decrypted and saved successfully")

        decryption_time = time.time() - start_time
        print(f"Decryption time for file: {decryption_time:.2f} s")

    except UnicodeDecodeError as e:
        print(f"Error decoding decrypted data: {e}")

def encrypt_message(message, key):
    start_time = time.time()
    iv = os.urandom(16)  # Generate IV
    cipher = Cipher(algorithms.AES(key), modes.CBC(iv))
    encryptor = cipher.encryptor()
    padder = padding.PKCS7(algorithms.AES.block_size).padder()
    padded_data = padder.update(message.encode()) + padder.finalize()
    encrypted_message = encryptor.update(padded_data) + encryptor.finalize()
    encryption_time = time.time() - start_time
    print(f"Encryption time for message: {encryption_time:.2f} s")

    return iv + encrypted_message  # Prepend IV to the encrypted message


def encrypt_file(file_path, key):
    start_time = time.time()
    if file_path.endswith('.docx'):
        content = read_word_file(file_path)
    elif file_path.endswith('.pdf'):
        content = read_pdf_file(file_path)
    else:
        with open(file_path, 'rb') as file:
            content = file.read()
    iv = os.urandom(16)  # Generate IV
    cipher = Cipher(algorithms.AES(key), modes.CBC(iv))
    encryptor = cipher.encryptor()
    padder = padding.PKCS7(algorithms.AES.block_size).padder()
    padded_data = padder.update(content) + padder.finalize()
    encrypted_content = encryptor.update(padded_data) + encryptor.finalize()

    encryption_time = time.time() - start_time
    print(f"Encryption time for file: {encryption_time:.2f} s")

    return iv + encrypted_content



def send_message_to_bob(data, key, message_listbox, recipient):
    server_address = ('127.0.0.2', 12345)
    client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    client_socket.connect(server_address)

    recipient_ip = client_socket.getpeername()[0]

    if recipient == "Nurse":
        allowed_ips = allowed_nurse_ips
    else:
        allowed_ips = allowed_doctor_ips

    if recipient_ip not in allowed_ips:
        print(f"Unauthorized access from IP: {recipient_ip}")
        client_socket.close()
        return

    mac_id = generate_mac_id_from_key(data, key)
    timestamp = time.ctime()

    if os.path.isfile(data):
        # Handle file sending
        encrypted_file = encrypt_file(data, key)

        client_socket.send(b"FILE:")  # Send "FILE:" protocol header
        client_socket.sendall(encrypted_file)

        print("File encrypted and sent successfully")
        message_listbox.insert(tk.END, f"{data} sent!!")
    else:
        # Handle message sending
        message_data = f"\nMAC id: {mac_id}\nTimestamp: {timestamp}\nMessage: {data}"
        message = encrypt_message(message_data, key)

        client_socket.send(b"MSG:")  # Send "MSG:" protocol header
        client_socket.sendall(message)

        print("Message encrypted and sent successfully")
        print(f"Message sent at time :{timestamp}")
        print(f"MAC ID: {mac_id}")

        # Display the sent message in the GUI
        message_listbox.insert(tk.END, f"[{timestamp}] You: {data}")

    client_socket.close()

def nurse_message(data, key, message_listbox, recipient):
    server_address = ('127.0.0.3', 12347)
    client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    client_socket.connect(server_address)

    recipient_ip = client_socket.getpeername()[0]

    if recipient == "Nurse":
        allowed_ips = allowed_nurse_ips
    else:
        allowed_ips = allowed_doctor_ips

    if recipient_ip not in allowed_ips:
        print(f"Unauthorized access from IP: {recipient_ip}")
        client_socket.close()
        return

    mac_id = generate_mac_id_from_key(data, key)
    timestamp = time.ctime()

    if os.path.isfile(data):
        # Handle file sending
        encrypted_file = encrypt_file(data, key)

        client_socket.send(b"FILE:")  # Send "FILE:" protocol header
        client_socket.sendall(encrypted_file)

        print("File encrypted and sent successfully")
    else:
        # Handle message sending
        message_data = f"\nMAC id: {mac_id}\nTimestamp: {timestamp}\nMessage: {data}"
        message = encrypt_message(message_data, key)

        client_socket.send(b"MSG:")  # Send "MSG:" protocol header
        client_socket.sendall(message)

        print("Message encrypted and sent successfully")
        print(f"Message sent at time :{timestamp}")
        print(f"MAC ID: {mac_id}")

        # Display the sent message in the GUI
        message_listbox.insert(tk.END, f"[{timestamp}] You: {data}")

    client_socket.close()

def send_message_to_doctor(message_entry, aes_key, message_listbox, recipient):
    message = message_entry.get()
    if message:
        # Handle regular message sending
        send_message_to_bob(message, aes_key, message_listbox, recipient)
        message_entry.delete(0, tk.END)

def send_message_to_nurse(message_entry, aes_key, message_listbox, recipient):
    message = message_entry.get()
    if message:
        # Handle regular message sending
        nurse_message(message, aes_key, message_listbox, recipient)
        message_entry.delete(0, tk.END)

#def upload_file_and_send(aes_key, message_entry, message_listbox):
#    root = tk.Tk()
#    file_path = filedialog.askopenfilename()  # Open file dialog to select a file
#    if file_path:
#        # Populate the message entry with the file path
#        filename = os.path.basename(file_path)
#
#        # Create a new window for recipient selection
#        send_to_window = tk.Toplevel(root)
#        send_to_window.title("Send To")
#
#        send_to_nurse_button = tk.Button(send_to_window, text="Send to Nurse",
#                                         command=partial(send_message_to_nurse, file_path, aes_key, message_listbox, "Nurse"))
#        send_to_nurse_button.pack(side=tk.TOP, pady=5)
#
#        send_to_doctor_button = tk.Button(send_to_window, text="Send to Doctor",
#                                          command=partial(send_message_to_doctor, file_path, aes_key, message_listbox, "Doctor"))
#        send_to_doctor_button.pack(side=tk.TOP, pady=5)
#
#        send_to_window.grab_set()  # Make this window modal (focus remains here)
#        send_to_window.wait_window()  # Wait for the window to close before continuing


def create_gui(aes_key):

    root = tk.Tk()
    root.title("Alice's Secure Message Exchange")

    message_frame = tk.Frame(root)
    message_frame.pack(side=tk.BOTTOM, fill=tk.X)

    message_entry = tk.Entry(message_frame, font=('Roboto', 12))  # Increase font size for emojis
    message_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=5)
    message_entry.focus_set()  # Set focus to the entry field

    message_log_frame = tk.Frame(root)
    message_log_frame.pack(expand=True, fill=tk.BOTH)

    message_log_scrollbar = tk.Scrollbar(message_log_frame)
    message_log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    message_listbox = tk.Listbox(message_log_frame, yscrollcommand=message_log_scrollbar.set, width=50, font=("Monaco", 12))
    message_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
    message_log_scrollbar.config(command=message_listbox.yview)

    nurse_button = tk.Button(message_frame, text="Send to Nurse", command=partial(send_message_to_nurse, message_entry, aes_key, message_listbox, "Nurse"))
    nurse_button.pack(side=tk.LEFT, padx=5, pady=5)

    doctor_button = tk.Button(message_frame, text="Send to Doctor Bob", command=partial(send_message_to_doctor, message_entry, aes_key, message_listbox, "Doctor"))
    doctor_button.pack(side=tk.RIGHT, padx=5, pady=5)

    #upload_button = tk.Button(message_frame, text="Upload File", command=partial(upload_file_and_send, aes_key, message_entry, message_listbox))
    #upload_button.pack(side=tk.LEFT, padx=5, pady=5)

    # Function to send message on pressing Enter key
    root.bind('<Return>', lambda event: send_message_to_doctor(message_entry, aes_key, message_listbox, "Doctor"))

    def start_server():
        server_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        server_address = ('127.0.0.1', 12346)
        server_socket.bind(server_address)
        server_socket.listen(3)

        try:
            while True:
                print("\n\nWaiting for a connection...")
                client_socket, client_address = server_socket.accept()
                print("Connected to", client_address)

                if client_address[0] in blocked_ips:
                    print(f"IP Address {client_address[0]} is blocked!")
                    client_socket.close()
                    continue  # Skip further processing for blocked IP addresses

                received_message = b""
                while True:
                    data = client_socket.recv(1024)
                    if not data:
                        break
                    received_message += data

                if received_message.startswith(b"MSG:"):
                    decrypted_message = decrypt_message(received_message[len(b"MSG:"):], aes_key)
                    if not decrypted_message:
                        print("Error decrypting message")
                        continue

                    message_lines = decrypted_message.split('\n')
                    if len(message_lines) > 1:
                        mac_id = message_lines[0].replace("MAC id: ", "")
                        timestamp_str = message_lines[1].replace("Timestamp: ", "")
                        message = message_lines[2].replace("Message: ", "")
                        calculated_mac_id = generate_mac_id_from_key(message, aes_key)
                        if calculated_mac_id == mac_id:
                            print("Validated MAC id... Message is authentic")
                        else:
                            print("Invalid MAC id..")
                        received_timestamp = datetime.strptime(timestamp_str, "%a %b %d %H:%M:%S %Y")
                        current_time = datetime.now()
                        time_window = 30
                        time_difference = current_time - received_timestamp
                        if time_difference.total_seconds() <= time_window:
                            print("Received message is within the time window.")
                            print("MAC id: ", mac_id)
                            print("Timestamp:", received_timestamp)
                            #print("Message:", message)
                            #message_logbox_entry = f"[{timestamp_str}]"
                            #message_listbox.insert(tk.END, message_logbox_entry)
                            message_listbox.see(tk.END)
                        else:
                            print("Received message is outside the time window. Ignoring.")

                        # Displaying all the necessary details in the GUI
                        #message_logbox_entry = f"MAC id: {mac_id}"
                        msg_str = message
                        message_listbox.insert(tk.END, f"[{received_timestamp}] Bob: {msg_str}")
                        message_listbox.see(tk.END)
                    else:
                        print("Cannot see mac id verification details")

                elif received_message.startswith(b"FILE:"):
                    file_content = received_message[len(b"FILE:"):]
                    decrypt_file(file_content, aes_key, os.path.join(os.path.expanduser('~'), 'Documents', 'file_from_alice.txt'))
                    print("File saved as 'file_from_bob.txt'.")
                    message_listbox.insert(tk.END, "File saved as 'file_from_bob.txt'.")
                    message_listbox.see(tk.END)

                client_socket.close()
        finally:
            server_socket.close()

    import threading
    threading.Thread(target=start_server).start()

    root.mainloop()


if __name__ == "__main__":
    create_gui(aes_key)
    delete_key_file()